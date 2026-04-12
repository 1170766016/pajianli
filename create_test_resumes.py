"""
生成测试用的 docx 和 pdf 简历文件
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_docx_resume():
    """生成 docx 格式测试简历"""
    doc = Document()

    # 标题 - 姓名
    title = doc.add_heading("陈思雨", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 联系方式
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run("📞 13912345678 | ✉️ chensiyu@example.com | 📍 上海")

    # 求职意向
    doc.add_heading("求职意向", level=2)
    doc.add_paragraph("Python 后端开发工程师")

    # 教育背景
    doc.add_heading("教育背景", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    cells = table.rows[0].cells
    cells[0].text = "2015-2019"
    cells[1].text = "浙江大学"
    cells[2].text = "软件工程"
    cells[3].text = "本科"

    # 工作经历
    doc.add_heading("工作经历", level=2)

    doc.add_paragraph("2022.03 - 至今  阿里巴巴  高级Python开发工程师（4年）", style="List Bullet")
    doc.add_paragraph("• 负责淘宝推荐系统后端核心模块的开发与维护")
    doc.add_paragraph("• 使用 FastAPI + PostgreSQL 设计高并发 RESTful API，QPS 达到 5000+")
    doc.add_paragraph("• 基于 Celery + RabbitMQ 实现异步任务调度系统")
    doc.add_paragraph("• 参与公司 AI 中台建设，对接大模型推理服务")

    doc.add_paragraph("2019.07 - 2022.02  网易  Python开发工程师（2.5年）", style="List Bullet")
    doc.add_paragraph("• 负责网易云音乐后端服务开发，使用 Django + MySQL")
    doc.add_paragraph("• 设计并实现用户画像系统，日活用户 2000 万+")
    doc.add_paragraph("• 使用 Redis 实现分布式缓存，接口响应时间降低 60%")

    # 技术栈
    doc.add_heading("技术栈", level=2)
    doc.add_paragraph("Python, FastAPI, Django, Flask, PostgreSQL, MySQL, Redis, RabbitMQ, Celery, Docker, Kubernetes, Git, Linux, AWS")

    # 项目经历
    doc.add_heading("项目经历", level=2)
    doc.add_paragraph("1. 智能内容审核平台（2023）")
    doc.add_paragraph("   基于 DeepSeek API + FastAPI 构建的内容审核系统，支持文本/图片多模态审核，日处理量 500 万+")
    doc.add_paragraph("2. 微服务网关系统（2022）")
    doc.add_paragraph("   基于 Kong + Python 插件的 API 网关，实现统一认证、限流、链路追踪")

    # 自我评价
    doc.add_heading("自我评价", level=2)
    doc.add_paragraph("6.5年 Python 后端开发经验，有大厂背景，熟悉高并发系统设计和大模型应用。具有良好的代码习惯和团队合作精神。")

    doc.save("resumes/陈思雨_Python开发.docx")
    print("✅ 已生成: resumes/陈思雨_Python开发.docx")


def create_pdf_resume():
    """生成 pdf 格式测试简历 (使用 reportlab)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import cm
        import os

        # 尝试注册中文字体
        font_name = "SimSun"
        font_paths = [
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/msyh.ttf",
        ]
        font_registered = False
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, fp))
                    font_registered = True
                    break
                except Exception:
                    continue

        if not font_registered:
            print("⚠️ 未找到中文字体，PDF 将使用英文内容")
            _create_pdf_english()
            return

        filepath = "resumes/刘海涛_全栈开发.pdf"
        c = canvas.Canvas(filepath, pagesize=A4)
        width, height = A4

        y = height - 2 * cm

        def draw_text(text, size=11, bold=False, indent=0):
            nonlocal y
            c.setFont(font_name, size)
            c.drawString(2 * cm + indent, y, text)
            y -= size * 1.5

        def draw_heading(text):
            nonlocal y
            y -= 0.3 * cm
            c.setFont(font_name, 14)
            c.drawString(2 * cm, y, text)
            y -= 0.2 * cm
            c.setStrokeColorRGB(0.2, 0.3, 0.6)
            c.setLineWidth(1)
            c.line(2 * cm, y, width - 2 * cm, y)
            y -= 0.5 * cm

        # 姓名
        c.setFont(font_name, 22)
        c.drawCentredString(width / 2, y, "刘海涛")
        y -= 1 * cm

        # 联系方式
        c.setFont(font_name, 10)
        c.drawCentredString(width / 2, y, "13687654321 | liuhaitao@example.com | 北京")
        y -= 0.8 * cm

        draw_heading("求职意向")
        draw_text("Python 后端开发工程师 / 全栈开发工程师")

        draw_heading("教育背景")
        draw_text("2013-2017  北京邮电大学  计算机科学与技术  本科")
        draw_text("2017-2020  清华大学  软件工程  硕士")

        draw_heading("工作经历")
        draw_text("2020.07 - 至今  美团  资深后端工程师（6年）")
        draw_text("• 负责美团外卖核心交易系统后端开发，使用 Python + Go 微服务架构", indent=0.5 * cm)
        draw_text("• 设计高可用订单系统，支撑日均 5000 万单交易", indent=0.5 * cm)
        draw_text("• 使用 Kafka + Flink 构建实时数据处理管道", indent=0.5 * cm)
        draw_text("• 主导 AI 智能调度系统的后端架构设计", indent=0.5 * cm)
        draw_text("• 带团队 8 人，负责技术方案评审和代码质量管理", indent=0.5 * cm)

        draw_heading("技术栈")
        draw_text("Python, Go, FastAPI, Django, Flask, PostgreSQL, MySQL, Redis")
        draw_text("Kafka, Docker, Kubernetes, AWS, Jenkins, Git, Linux")

        draw_heading("项目经历")
        draw_text("1. 智能配送调度系统（2023）")
        draw_text("  基于深度强化学习 + Python 后端的外卖配送路径优化系统", indent=0.5 * cm)
        draw_text("  配送效率提升 15%，骑手满意度提升 20%", indent=0.5 * cm)

        draw_text("2. 商家数据分析平台（2022）")
        draw_text("  基于 FastAPI + React 的全栈数据分析平台", indent=0.5 * cm)
        draw_text("  支持实时经营数据看板，服务 100 万+ 商家", indent=0.5 * cm)

        draw_heading("自我评价")
        draw_text("硕士学历，6年后端开发经验，具备全栈能力。")
        draw_text("有大型互联网公司核心业务开发经验，擅长高并发系统设计。")
        draw_text("有团队管理经验，有开源项目贡献。")

        c.save()
        print(f"✅ 已生成: {filepath}")

    except ImportError:
        print("⚠️ 未安装 reportlab，使用 fpdf2 生成 PDF")
        _create_pdf_with_fpdf()


def _create_pdf_english():
    """如果没有中文字体，生成英文 PDF"""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import cm

    filepath = "resumes/liuhaitao_fullstack.pdf"
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(width / 2, y, "Liu Haitao")
    y -= 1 * cm

    c.setFont("Helvetica", 10)
    c.drawCentredString(width / 2, y, "13687654321 | liuhaitao@example.com | Beijing")
    y -= 1 * cm

    lines = [
        ("Helvetica-Bold", 14, "Education"),
        ("Helvetica", 11, "2013-2017  BUPT  Computer Science  Bachelor"),
        ("Helvetica", 11, "2017-2020  Tsinghua University  Software Engineering  Master"),
        ("Helvetica-Bold", 14, ""),
        ("Helvetica-Bold", 14, "Work Experience"),
        ("Helvetica", 11, "2020.07 - Present  Meituan  Senior Backend Engineer (6 years)"),
        ("Helvetica", 11, "  - Core trading system development with Python + Go microservices"),
        ("Helvetica", 11, "  - Designed high-availability order system, 50M+ daily orders"),
        ("Helvetica", 11, "  - Built real-time data pipeline with Kafka + Flink"),
        ("Helvetica", 11, "  - Led AI-powered delivery scheduling system"),
        ("Helvetica", 11, "  - Managed team of 8 engineers"),
        ("Helvetica-Bold", 14, ""),
        ("Helvetica-Bold", 14, "Skills"),
        ("Helvetica", 11, "Python, Go, FastAPI, Django, Flask, PostgreSQL, MySQL, Redis"),
        ("Helvetica", 11, "Kafka, Docker, Kubernetes, AWS, Jenkins, Git, Linux"),
        ("Helvetica-Bold", 14, ""),
        ("Helvetica-Bold", 14, "Summary"),
        ("Helvetica", 11, "Master degree, 6 years backend experience, full-stack capable."),
        ("Helvetica", 11, "Expert in high-concurrency system design with open source contributions."),
    ]

    for font, size, text in lines:
        c.setFont(font, size)
        c.drawString(2 * cm, y, text)
        y -= size * 1.5

    c.save()
    print(f"✅ Generated: {filepath}")


def _create_pdf_with_fpdf():
    """使用 fpdf2 作为备选方案"""
    try:
        from fpdf import FPDF
    except ImportError:
        # 最后的备选方案：直接写一个最简 PDF
        _create_minimal_pdf()
        return

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("msyh", "", "C:/Windows/Fonts/msyh.ttc", uni=True)
    pdf.set_font("msyh", size=22)
    pdf.cell(0, 15, "刘海涛", align="C", ln=True)
    pdf.set_font("msyh", size=10)
    pdf.cell(0, 8, "13687654321 | liuhaitao@example.com | 北京", align="C", ln=True)

    sections = [
        ("求职意向", ["Python 后端开发工程师 / 全栈开发工程师"]),
        ("教育背景", ["2017-2020  清华大学  软件工程  硕士", "2013-2017  北京邮电大学  计算机科学与技术  本科"]),
        ("工作经历", [
            "2020.07 - 至今  美团  资深后端工程师（6年）",
            "• 负责美团外卖核心交易系统后端开发",
            "• 设计高可用订单系统，支撑日均 5000 万单交易",
            "• 使用 Kafka + Flink 构建实时数据处理管道",
            "• 主导 AI 智能调度系统的后端架构设计",
        ]),
        ("技术栈", ["Python, Go, FastAPI, Django, Flask, PostgreSQL, MySQL, Redis, Kafka, Docker, K8s"]),
        ("自我评价", ["硕士学历，6年后端开发经验，有大厂核心业务开发和团队管理经验。"]),
    ]

    for title, items in sections:
        pdf.ln(3)
        pdf.set_font("msyh", size=14)
        pdf.cell(0, 10, title, ln=True)
        pdf.set_font("msyh", size=11)
        for item in items:
            pdf.cell(0, 7, item, ln=True)

    pdf.output("resumes/刘海涛_全栈开发.pdf")
    print("✅ 已生成: resumes/刘海涛_全栈开发.pdf")


def _create_minimal_pdf():
    """最简方案：纯文本写入 pdf（用于没有字体库的环境）"""
    # 写一个简单的文本文件作为替代
    content = """刘海涛

联系方式：13687654321 | liuhaitao@example.com | 北京

求职意向：Python 后端开发工程师 / 全栈开发工程师

教育背景：
2017-2020  清华大学  软件工程  硕士
2013-2017  北京邮电大学  计算机科学与技术  本科

工作经历：
2020.07 - 至今  美团  资深后端工程师（6年）
- 负责美团外卖核心交易系统后端开发，使用 Python + Go 微服务架构
- 设计高可用订单系统，支撑日均 5000 万单交易
- 使用 Kafka + Flink 构建实时数据处理管道
- 主导 AI 智能调度系统的后端架构设计
- 带团队 8 人，负责技术方案评审和代码质量管理

技术栈：
Python, Go, FastAPI, Django, Flask, PostgreSQL, MySQL, Redis, Kafka, Docker, Kubernetes

项目经历：
1. 智能配送调度系统（2023） - 基于深度强化学习 + Python 后端的外卖配送路径优化
2. 商家数据分析平台（2022） - 基于 FastAPI + React 的全栈数据分析平台

自我评价：
硕士学历，6年后端开发经验，具备全栈能力。有团队管理经验和开源项目贡献。
"""
    # 如果无法生成 PDF，保存为 txt 格式
    with open("resumes/刘海涛_全栈开发.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("⚠️ PDF 库不可用，已生成 txt 替代: resumes/刘海涛_全栈开发.txt")


if __name__ == "__main__":
    create_docx_resume()
    create_pdf_resume()
