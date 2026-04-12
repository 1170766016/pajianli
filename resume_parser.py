"""
简历解析模块 - 支持 HTML / PDF / Word / TXT 格式
核心思路：提取简历的纯文本，交给 LLM 做智能理解
"""
import os
import re
import sys
from pathlib import Path

# Windows 终端 UTF-8 编码支持
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

from models import ResumeData


def parse_resume(file_path: str) -> ResumeData:
    """
    解析单份简历文件，自动识别格式
    
    Args:
        file_path: 简历文件的绝对路径
        
    Returns:
        ResumeData 对象
    """
    file_path = os.path.abspath(file_path)
    ext = Path(file_path).suffix.lower()
    file_name = Path(file_path).name

    # 根据后缀分发到对应解析器
    parsers = {
        ".html": _parse_html,
        ".htm": _parse_html,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".doc": _parse_doc,
        ".txt": _parse_txt,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {list(parsers.keys())}")

    raw_text = parser(file_path)

    if not raw_text or len(raw_text.strip()) < 20:
        raise ValueError(f"简历内容过少或解析失败: {file_name}")

    # 构建简历数据对象
    resume = ResumeData(
        file_path=file_path,
        file_name=file_name,
        raw_text=raw_text,
    )

    # 尝试用正则提取基本信息（辅助，主要靠 LLM）
    _extract_basic_info(resume)

    return resume


def parse_all_resumes(directory: str) -> list[ResumeData]:
    """
    扫描目录下所有简历文件并解析
    
    Args:
        directory: 简历文件夹路径
        
    Returns:
        ResumeData 列表
    """
    from config import SUPPORTED_FORMATS

    resumes = []
    errors = []

    if not os.path.isdir(directory):
        raise FileNotFoundError(f"简历目录不存在: {directory}")

    files = sorted(Path(directory).iterdir())
    resume_files = [f for f in files if f.suffix.lower() in SUPPORTED_FORMATS]

    if not resume_files:
        raise FileNotFoundError(
            f"在 {directory} 中未找到简历文件\n"
            f"支持的格式: {SUPPORTED_FORMATS}\n"
            f"请将简历文件放入该目录"
        )

    for file_path in resume_files:
        try:
            resume = parse_resume(str(file_path))
            resumes.append(resume)
        except Exception as e:
            errors.append(f"  ✗ {file_path.name}: {e}")

    if errors:
        print(f"\n⚠️  以下 {len(errors)} 份简历解析失败:")
        for err in errors:
            print(err)

    return resumes


# ============================================================
# 各格式解析器
# ============================================================

def _parse_html(file_path: str) -> str:
    """解析 HTML 格式简历（前程无忧下载的通常是 HTML）"""
    from bs4 import BeautifulSoup

    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        html = f.read()

    soup = BeautifulSoup(html, "lxml")

    # 移除 script 和 style 标签
    for tag in soup(["script", "style", "meta", "link"]):
        tag.decompose()

    # 提取文本
    text = soup.get_text(separator="\n", strip=True)

    # 清理多余空行
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _parse_pdf(file_path: str) -> str:
    """解析 PDF 格式简历"""
    import pdfplumber

    texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)

    return "\n\n".join(texts).strip()


def _parse_docx(file_path: str) -> str:
    """解析 Word (.docx) 格式简历"""
    from docx import Document

    doc = Document(file_path)
    texts = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())

    # 也提取表格中的内容（简历常用表格布局）
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)

    return "\n".join(texts).strip()


def _parse_doc(file_path: str) -> str:
    """
    解析旧版 Word (.doc) 格式简历
    注意：.doc 格式需要系统安装 antiword 或通过 LibreOffice 转换
    如果都没有，回退到提取二进制中的文本
    """
    import subprocess
    import shutil

    # 方法1：尝试用 antiword
    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass

    # 方法2：尝试用 LibreOffice 转换
    if shutil.which("libreoffice") or shutil.which("soffice"):
        try:
            cmd = shutil.which("libreoffice") or shutil.which("soffice")
            output_dir = os.path.dirname(file_path)
            subprocess.run(
                [cmd, "--headless", "--convert-to", "txt", "--outdir", output_dir, file_path],
                capture_output=True, timeout=60
            )
            txt_path = os.path.splitext(file_path)[0] + ".txt"
            if os.path.exists(txt_path):
                with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                os.remove(txt_path)  # 清理临时文件
                return text.strip()
        except Exception:
            pass

    # 方法3：暴力提取二进制中的中文和英文文本
    with open(file_path, "rb") as f:
        raw = f.read()
    # 尝试提取 UTF-8 或 GBK 编码的文本
    for encoding in ["utf-8", "gbk", "gb2312"]:
        try:
            text = raw.decode(encoding, errors="ignore")
            # 只保留可读字符
            text = re.sub(r'[^\u4e00-\u9fff\u0020-\u007ea-zA-Z0-9@.，。、；：""''（）\n]', '', text)
            if len(text) > 50:
                return text.strip()
        except Exception:
            continue

    raise ValueError(f".doc 文件解析失败，建议转换为 .docx 或 .pdf 格式: {file_path}")


def _parse_txt(file_path: str) -> str:
    """解析纯文本格式简历"""
    for encoding in ["utf-8", "gbk", "gb2312", "utf-16"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read().strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


# ============================================================
# 辅助信息提取（正则，辅助用）
# ============================================================

def _extract_basic_info(resume: ResumeData):
    """用正则从文本中提取基本信息，作为辅助"""
    text = resume.raw_text

    # 提取手机号
    phone_match = re.search(r'1[3-9]\d{9}', text)
    if phone_match:
        resume.phone = phone_match.group()

    # 提取邮箱
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.]+', text)
    if email_match:
        resume.email = email_match.group()

    # 提取姓名（通常在简历开头，取第一行非空文本作为候选）
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        first_line = lines[0]
        # 如果第一行很短（2-4个字），大概率是姓名
        if 2 <= len(first_line) <= 4 and re.match(r'^[\u4e00-\u9fff]+$', first_line):
            resume.name = first_line

    # 提取工作年限
    years_match = re.search(r'(\d+)\s*[年].*?[工经].*?[作验]', text)
    if years_match:
        resume.work_years = f"{years_match.group(1)}年"

    # 提取学历关键词
    edu_keywords = ["博士", "硕士", "研究生", "本科", "大专", "专科", "大学"]
    for keyword in edu_keywords:
        if keyword in text:
            resume.education = keyword
            break
